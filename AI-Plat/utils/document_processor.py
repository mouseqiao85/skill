"""
Document Processor for AI-Plat
Converts various document formats to text for AI processing
Supports .docx, .pdf, .txt, and other common formats
"""

import os
import sys
from pathlib import Path
import json
from typing import Optional, Dict, Any


def process_document(input_path: str, output_path: Optional[str] = None) -> Optional[str]:
    """
    Process a document and convert it to text format
    Supports .docx, .pdf, .txt, and other formats
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"Error: File {input_path} does not exist")
        return None
    
    if output_path is None:
        output_path = input_path.with_suffix('.txt')
    else:
        output_path = Path(output_path)
    
    # Determine the appropriate converter based on file extension
    if input_path.suffix.lower() == '.docx':
        return _process_docx(input_path, output_path)
    elif input_path.suffix.lower() == '.pdf':
        return _process_pdf(input_path, output_path)
    elif input_path.suffix.lower() in ['.txt', '.md', '.rst']:
        return _process_text(input_path, output_path)
    else:
        print(f"Unsupported file format: {input_path.suffix}")
        return None


def _process_docx(docx_path: Path, output_path: Path) -> Optional[str]:
    """Process a .docx file"""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(f"[TABLE CELL] {cell.text}")
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Header
            header_paragraphs = []
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    header_paragraphs.append(paragraph.text)
            if header_paragraphs:
                full_text.extend(["[HEADER START]", *header_paragraphs, "[HEADER END]"])
            
            # Footer
            footer_paragraphs = []
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    footer_paragraphs.append(paragraph.text)
            if footer_paragraphs:
                full_text.extend(["[FOOTER START]", *footer_paragraphs, "[FOOTER END]"])
        
        text_content = '\n'.join(full_text)
        
        # Save to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        print(f"Successfully converted {docx_path} to {output_path}")
        print(f"Extracted {len(full_text)} text elements")
        
        return text_content
        
    except ImportError:
        print("python-docx library not found. Please ensure it's installed: pip install python-docx")
        return None
    except Exception as e:
        print(f"Error processing .docx file: {e}")
        return None


def _process_pdf(pdf_path: Path, output_path: Path) -> Optional[str]:
    """Process a .pdf file"""
    try:
        import PyPDF2
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_pages = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_pages.append(f"[PAGE {page_num + 1}]")
                    text_pages.append(text)
                    text_pages.append("")  # Empty line between pages
        
        text_content = '\n'.join(text_pages)
        
        # Save to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        print(f"Successfully converted {pdf_path} to {output_path}")
        print(f"Extracted {len(pdf_reader.pages)} pages")
        
        return text_content
        
    except ImportError:
        print("PyPDF2 library not found. Please install it: pip install PyPDF2")
        return None
    except Exception as e:
        print(f"Error processing .pdf file: {e}")
        return None


def _process_text(txt_path: Path, output_path: Path) -> Optional[str]:
    """Process a text-based file (.txt, .md, .rst)"""
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        # Just copy to output if different
        if txt_path != output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            print(f"Successfully copied {txt_path} to {output_path}")
        else:
            print(f"Processed {txt_path}")
        
        return text_content
        
    except Exception as e:
        print(f"Error processing text file: {e}")
        return None


def split_long_text(text: str, max_chars: int = 2000) -> list:
    """
    Split long text into smaller chunks for processing
    """
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    for i in range(0, len(text), max_chars):
        chunk = text[i:i + max_chars]
        # Try to break at sentence boundary if possible
        if i + max_chars < len(text):
            # Find the last period in the chunk
            last_period = chunk.rfind('.')
            if last_period > max_chars * 0.8:  # If period is in the last 20%
                chunk = chunk[:last_period + 1]
        
        chunks.append(chunk.strip())
    
    return chunks


def main():
    """Main function to process documents"""
    if len(sys.argv) < 2:
        print("Document Processor for AI-Plat")
        print("Converts various document formats to text for AI processing")
        print("\nUsage: python document_processor.py <input_file> [output_file]")
        print("\nSupported formats: .docx, .pdf, .txt, .md, .rst")
        print("\nExample: python document_processor.py document.docx output.txt")
        return
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    print(f"Processing document: {input_path}")
    
    result = process_document(input_path, output_path)
    
    if result:
        print(f"\nDocument processed successfully!")
        print(f"Length: {len(result)} characters")
        
        # Show preview if the text is not too long
        if len(result) < 1000:
            print("\nContent preview:")
            print("-" * 50)
            print(result)
            print("-" * 50)
        else:
            print("\nContent preview (first 500 characters):")
            print("-" * 50)
            print(result[:500] + "...")
            print("-" * 50)
    else:
        print("Failed to process the document")


if __name__ == "__main__":
    main()