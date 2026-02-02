"""
Utility script to convert .docx files to text format
This allows the AI to process Word documents by converting them to text first
"""

def docx_to_text(docx_path):
    """
    Convert a .docx file to plain text
    Requires python-docx library: pip install python-docx
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except ImportError:
        print("python-docx library not found. Please install it using: pip install python-docx")
        return None
    except Exception as e:
        print(f"Error reading document: {e}")
        return None

def docx_to_markdown(docx_path):
    """
    Convert a .docx file to Markdown format (more structured than plain text)
    """
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re
        
        doc = Document(docx_path)
        md_lines = []
        
        for element in doc.paragraphs:
            text = element.text.strip()
            if not text:
                continue
                
            # Handle headings based on style
            if element.style.name.startswith('Heading'):
                level = element.style.name[-1]  # Get the number from Heading1, Heading2, etc.
                if level.isdigit():
                    md_lines.append(f"{'#' * int(level)} {text}")
                else:
                    md_lines.append(f"# {text}")
            else:
                # Handle regular paragraphs
                md_lines.append(text)
                
                # Check for alignment
                if element.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    md_lines.append("")  # Add spacing for centered content
        
        # Process tables separately
        for table in doc.tables:
            if len(table.rows) > 0:
                md_lines.append("")  # Add blank line before table
                # Add header row
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                md_lines.append("| " + " | ".join(header_cells) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                
                # Add data rows
                for row in table.rows[1:]:
                    data_cells = [cell.text.strip() for cell in row.cells]
                    md_lines.append("| " + " | ".join(data_cells) + " |")
                md_lines.append("")  # Add blank line after table
        
        return '\n'.join(md_lines)
    except ImportError:
        print("python-docx library not found. Please install it using: pip install python-docx")
        return None
    except Exception as e:
        print(f"Error reading document: {e}")
        return None

def save_text_file(content, output_path):
    """
    Save text content to a file
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"Content saved to {output_path}")
        return True
    except Exception as e:
        print(f"Error saving file: {e}")
        return False

def main():
    """
    Main function to convert docx to text/markdown
    """
    import sys
    import os
    
    if len(sys.argv) < 2:
        print("Usage: python docx_to_text.py <input.docx> [output.txt]")
        print("If no output file is specified, it will create input.txt")
        return
    
    input_path = sys.argv[1]
    
    if not os.path.exists(input_path):
        print(f"File {input_path} does not exist")
        return
    
    if len(sys.argv) > 2:
        output_path = sys.argv[2]
    else:
        # Create output path by changing extension to .txt
        output_path = os.path.splitext(input_path)[0] + '.txt'
    
    print(f"Converting {input_path} to {output_path}...")
    
    # Try to convert to text
    text_content = docx_to_text(input_path)
    
    if text_content:
        success = save_text_file(text_content, output_path)
        if success:
            print("Conversion completed successfully!")
            print("\nFirst 500 characters of converted content:")
            print("-" * 50)
            print(text_content[:500] + ("..." if len(text_content) > 500 else ""))
            print("-" * 50)
        else:
            print("Failed to save the converted content")
    else:
        print("Failed to convert the document")

if __name__ == "__main__":
    main()