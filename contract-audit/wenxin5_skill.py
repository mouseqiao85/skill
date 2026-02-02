#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Contract Audit Skill Implementation
Based on the contract-audit skill specification for requirements analysis
"""

import pandas as pd
import os
import sys
from typing import Dict, List, Tuple, Optional
import re

# 添加对docx的支持
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with 'pip install python-docx' for DOC output support.")


class ContractAuditSkill:
    """
    Contract Audit Skill for requirements analysis and risk assessment
    """
    
    def __init__(self, traindata_path: str):
        """
        Initialize the skill with training data
        
        Args:
            traindata_path: Path to the traindata directory
        """
        self.traindata_path = traindata_path
        self.product_features = self._load_product_features()
        self.function_points = self._load_function_points()
        self.risk_examples = self._load_risk_examples()
    
    def _load_product_features(self) -> pd.DataFrame:
        """Load product feature list from traindata/list/"""
        try:
            list_dir = os.path.join(self.traindata_path, "list")
            excel_files = [f for f in os.listdir(list_dir) if f.endswith('.xlsx')]
            if excel_files:
                file_path = os.path.join(list_dir, excel_files[0])
                df = pd.read_excel(file_path)
                
                # Handle encoded column names
                actual_columns = df.columns.tolist()
                columns_mapping = {}
                if len(actual_columns) >= 7:
                    columns_mapping = {
                        actual_columns[0]: '序号',
                        actual_columns[1]: '模块名称', 
                        actual_columns[2]: '一级菜单',
                        actual_columns[3]: '二级菜单',
                        actual_columns[4]: '三级菜单',
                        actual_columns[5]: '功能清单',
                        actual_columns[6]: '功能描述(具体需要实现的功能)',
                        actual_columns[7]: '现有产品SOW及功能实现路径和客户指标功能',
                        actual_columns[8]: '是否关闭',
                        actual_columns[9]: '备注或其他数据表示'
                    }
                
                df = df.rename(columns=columns_mapping)
                # Only keep meaningful rows
                df = df[df['模块名称'].notna() | df['功能清单'].notna()]
                return df
        except Exception as e:
            print(f"Error loading product features: {e}")
        
        # Return empty dataframe if loading fails
        return pd.DataFrame(columns=['模块名称', '功能清单', '功能描述(具体需要实现的功能)'])
    
    def _load_function_points(self) -> List[Dict]:
        """Load function point evaluation data"""
        try:
            eval_dir = os.path.join(self.traindata_path, "功能点工作量评估")
            excel_files = [f for f in os.listdir(eval_dir) if f.endswith('.xlsx')]
            
            function_points = []
            for file in excel_files:
                file_path = os.path.join(eval_dir, file)
                df = pd.read_excel(file_path)
                # Extract function point data
                for _, row in df.iterrows():
                    function_points.append(row.to_dict())
            return function_points
        except Exception as e:
            print(f"Error loading function points: {e}")
            return []
    
    def _load_risk_examples(self) -> List[Dict]:
        """Load risk classification examples"""
        try:
            risk_dir = os.path.join(self.traindata_path, "四级风险清单")
            excel_files = [f for f in os.listdir(risk_dir) if f.endswith('.xlsx')]
            
            risks = []
            for file in excel_files:
                file_path = os.path.join(risk_dir, file)
                df = pd.read_excel(file_path)
                for _, row in df.iterrows():
                    risks.append(row.to_dict())
            return risks
        except Exception as e:
            print(f"Error loading risk examples: {e}")
            return []
    
    def analyze_contract(self, contract_content: str) -> Dict:
        """
        Analyze contract content and generate audit report
        
        Args:
            contract_content: Text content of the contract to analyze
            
        Returns:
            Dictionary containing the analysis results
        """
        # Step 1: Requirements Decomposition
        requirements = self._extract_requirements(contract_content)
        
        # Step 2: Product Matching Analysis
        matching_result = self._analyze_product_matching(requirements)
        
        # Step 3: Technical Risk Identification
        tech_risks = self._identify_technical_risks(requirements, matching_result)
        
        # Step 4: Non-functional Requirements Analysis
        non_func_reqs = self._analyze_non_functional_requirements(contract_content)
        
        # Step 5: Workload Estimation
        workload_estimate = self._estimate_workload(matching_result)
        
        # Step 6: Risk Consolidation
        consolidated_risks = self._consolidate_risks(matching_result, non_func_reqs, tech_risks)
        
        # Step 7: Four-Level Risk Classification
        four_level_risks = self._generate_four_level_risks(consolidated_risks)
        
        return {
            'requirements': requirements,
            'matching_analysis': matching_result,
            'technical_risks': tech_risks,
            'non_functional_analysis': non_func_reqs,
            'workload_estimate': workload_estimate,
            'consolidated_risks': consolidated_risks,
            'four_level_risks': four_level_risks
        }
    
    def _extract_requirements(self, content: str) -> Dict:
        """Extract functional and non-functional requirements from contract content"""
        # Define patterns for different types of requirements
        functional_patterns = [
            r"(\w+)功能",
            r"(\w+)系统",
            r"(\w+)平台",
            r"(\w+)服务",
            r"(\w+)管理",
            r"(\w+)处理",
            r"(\w+)分析",
            r"(\w+)监控",
            r"(\w+)接口",
            r"(\w+)集成"
        ]
        
        performance_patterns = [
            r"响应时间.*?(\d+)",
            r"并发.*?(\d+)",
            r"吞吐量.*?(\d+)",
            r"可用性.*?(\d+)%"
        ]
        
        security_patterns = [
            r"安全",
            r"加密",
            r"权限",
            r"认证",
            r"审计",
            r"隐私"
        ]
        
        # Extract functional requirements
        functional_reqs = []
        for pattern in functional_patterns:
            matches = re.findall(pattern, content)
            functional_reqs.extend(matches)
        
        # Extract performance requirements
        performance_reqs = []
        for pattern in performance_patterns:
            matches = re.findall(pattern, content)
            performance_reqs.extend(matches)
        
        # Extract security requirements
        security_reqs = []
        for pattern in security_patterns:
            if re.search(pattern, content):
                security_reqs.append(pattern)
        
        return {
            'functional': list(set(functional_reqs)),
            'performance': list(set(performance_reqs)),
            'security': list(set(security_reqs)),
            'integration': [],
            'compliance': []
        }
    
    def _analyze_product_matching(self, requirements: Dict) -> Dict:
        """Analyze how well requirements match existing products"""
        # Get unique functional requirements
        func_reqs = set(requirements.get('functional', []))
        
        # Match against product features
        matched_features = []
        unmatched_features = []
        
        if not self.product_features.empty:
            for req in func_reqs:
                # Check if requirement matches any product feature
                matches = self.product_features[
                    self.product_features['功能清单'].str.contains(req, na=False, case=False) |
                    self.product_features['模块名称'].str.contains(req, na=False, case=False) |
                    self.product_features['功能描述(具体需要实现的功能)'].str.contains(req, na=False, case=False)
                ]
                
                if not matches.empty:
                    matched_features.append({
                        'requirement': req,
                        'matched_features': matches[['模块名称', '功能清单']].to_dict('records')
                    })
                else:
                    unmatched_features.append(req)
        
        match_percentage = len(matched_features) / max(len(func_reqs), 1) * 100
        
        return {
            'match_percentage': match_percentage,
            'matched_features': matched_features,
            'unmatched_features': unmatched_features,
            'summary': f"产品功能匹配度约为 {int(match_percentage)}%"
        }
    
    def _identify_technical_risks(self, requirements: Dict, matching_result: Dict) -> List[Dict]:
        """Identify technical risks based on requirements and matching"""
        risks = []
        
        # Identify unsatisfied points
        unmatched = matching_result.get('unmatched_features', [])
        if unmatched:
            risks.append({
                'type': '功能不满足',
                'description': f"以下需求无法通过现有产品满足: {', '.join(unmatched[:5])}",
                'severity': 'high'
            })
        
        # Identify unclear requirements
        if not requirements.get('performance'):
            risks.append({
                'type': '性能指标不明确',
                'description': "合同中未明确具体的性能指标要求",
                'severity': 'medium'
            })
        
        if not requirements.get('security'):
            risks.append({
                'type': '安全要求不明确',
                'description': "合同中未明确具体的安全要求",
                'severity': 'high'
            })
        
        return risks
    
    def _analyze_non_functional_requirements(self, content: str) -> Dict:
        """Analyze non-functional requirements"""
        non_func = {
            'performance': [],
            'security': [],
            'availability': [],
            'scalability': [],
            'compliance': []
        }
        
        # Extract various non-functional requirements
        if '响应时间' in content:
            non_func['performance'].append('响应时间要求')
        
        if '安全' in content or '加密' in content:
            non_func['security'].append('安全要求')
        
        if '可用性' in content or '99%' in content:
            non_func['availability'].append('高可用要求')
        
        if '扩展' in content or '扩容' in content:
            non_func['scalability'].append('可扩展性要求')
        
        if '合规' in content or '监管' in content:
            non_func['compliance'].append('合规性要求')
        
        return non_func
    
    def _estimate_workload(self, matching_result: Dict) -> Dict:
        """Estimate workload based on matching analysis"""
        # Use function point data to estimate effort
        unmatched_count = len(matching_result.get('unmatched_features', []))
        
        # Base estimate: 8 person-days per function point based on historical data
        base_effort_per_fp = 8.0
        
        # Adjust for complexity
        complexity_factor = 1.2 if unmatched_count > 5 else 1.0
        
        estimated_effort = unmatched_count * base_effort_per_fp * complexity_factor
        
        return {
            'unmatched_items_count': unmatched_count,
            'estimated_person_days': round(estimated_effort, 2),
            'complexity_factor': complexity_factor,
            'base_effort_per_item': base_effort_per_fp
        }
    
    def _consolidate_risks(self, matching_result: Dict, non_func_reqs: Dict, tech_risks: List[Dict]) -> List[Dict]:
        """Consolidate all risks into a unified view"""
        consolidated = []
        
        # Add product capability risks
        match_pct = matching_result.get('match_percentage', 0)
        if match_pct < 80:
            consolidated.append({
                'category': '产品能力风险',
                'description': f"产品功能匹配度仅为{match_pct:.0f}%，存在较高定制开发风险",
                'severity': 'high' if match_pct < 70 else 'medium',
                'type': 'capability_gap'
            })
        
        # Add non-functional requirement risks
        for category, items in non_func_reqs.items():
            if items and category != 'compliance':
                consolidated.append({
                    'category': f"非功能性需求-{category}",
                    'description': f"{category}类非功能需求较多，可能存在实现风险",
                    'severity': 'medium',
                    'type': 'non_functional'
                })
        
        # Add technical risks
        for risk in tech_risks:
            consolidated.append({
                'category': f"技术风险-{risk['type']}",
                'description': risk['description'],
                'severity': risk['severity'],
                'type': 'technical'
            })
        
        return consolidated
    
    def _generate_four_level_risks(self, consolidated_risks: List[Dict]) -> List[Dict]:
        """Generate four-level risk classification"""
        four_level_risks = []
        
        for i, risk in enumerate(consolidated_risks):
            risk_entry = {
                '风险事项': risk['description'][:50] + "..." if len(risk['description']) > 50 else risk['description'],
                '风险等级': risk['severity'].upper(),
                '风险状态': '待处理',
                '预计发生时间': '项目执行期间',
                '风险责任人': '项目经理',
                '风险影响(成本)': '10-20%' if risk['severity'] == 'high' else '5-10%',
                '风险影响(进度)': '1-2周' if risk['severity'] == 'high' else '3-5天',
                '风险影响(质量)': '降级' if risk['severity'] == 'high' else '轻微影响',
                '风险触发阈值': risk['category'],
                '风险描述': risk['description'],
                '风险展朌': '可能导致项目延期或成本超支' if risk['severity'] == 'high' else '可能需要额外资源投入',
                '应对措施': self._get_default_mitigation(risk['type']),
                '风险跟踪': '定期跟踪'
            }
            four_level_risks.append(risk_entry)
        
        return four_level_risks
    
    def _get_default_mitigation(self, risk_type: str) -> str:
        """Get default mitigation strategy based on risk type"""
        mitigations = {
            'capability_gap': '1.加强需求调研 2.制定详细技术方案 3.预留缓冲时间',
            'non_functional': '1.明确非功能指标 2.制定测试计划 3.性能优化预案',
            'technical': '1.技术预研 2.原型验证 3.分阶段实施'
        }
        return mitigations.get(risk_type, '1.制定应对计划 2.定期评估 3.及时调整')
    
    def generate_report_md(self, analysis_results: Dict) -> str:
        """Generate formatted markdown report from analysis results"""
        report_parts = [
            "# 合同需求评估报告\n",
            "## 一、需求拆解清单\n",
            "### 1.1 功能需求\n",
            "根据合同内容，识别出以下功能需求：\n"
        ]
        
        # Add functional requirements
        func_reqs = analysis_results['requirements']['functional']
        for req in func_reqs[:10]:  # Limit to first 10 for readability
            report_parts.append(f"- **{req}功能**: 基于合同条款识别\n")
        
        report_parts.extend([
            "\n### 1.2 集成需求\n",
            "- 与现有系统的集成对接需求\n",
            "- 支持多种技术架构\n",
            "\n### 1.3 非功能需求\n",
            "\n#### 性能要求\n",
            "- 系统性能指标需进一步明确\n",
            "\n#### 安全要求\n",
            "- 数据安全和权限管理要求\n",
            "\n#### 高可用要求\n",
            "- 系统可用性要求\n",
            "\n## 二、产品功能匹配度分析\n",
            f"\n### 2.1 与现有产品的匹配情况\n",
            f"根据产品功能清单分析，合同中产品功能匹配度约为 **{analysis_results['matching_analysis']['match_percentage']:.0f}%**\n",
            "\n### 2.2 不满足的需求\n"
        ])
        
        # Add detailed unmatched features list
        unmatched = analysis_results['matching_analysis']['unmatched_features']
        if unmatched:
            report_parts.append("\n#### 完全不满足的需求（与产品功能清单匹配不满足内容清单）\n")
            for i, item in enumerate(unmatched, 1):
                report_parts.append(f"{i}. **{item}**: 现有产品暂不支持此功能，需要定制开发\n")
        
        # Add matched features for transparency
        matched_details = analysis_results['matching_analysis']['matched_features']
        if matched_details:
            report_parts.append("\n#### 满足的需求（与产品功能清单匹配内容）\n")
            for match_info in matched_details:
                req = match_info['requirement']
                report_parts.append(f"- **{req}**: 与以下产品功能匹配\n")
                for feature in match_info['matched_features'][:2]:  # Limit to first 2 for readability
                    module_name = feature.get('模块名称', '未知')
                    feature_name = feature.get('功能清单', '未知')
                    report_parts.append(f"  - 模块: {module_name}, 功能: {feature_name}\n")
        
        report_parts.extend([
            "\n## 三、技术风险识别\n",
            "\n### 3.1 不满足的点\n"
        ])
        
        tech_risks = analysis_results['technical_risks']
        for risk in tech_risks:
            report_parts.append(f"- **{risk['type']}**: {risk['description']}\n")
        
        report_parts.extend([
            "\n## 四、非功能需求拆解\n",
            "\n### 4.1 性能需求敞口\n",
            "- 性统性能指标需进一步明确\n",
            "\n### 4.2 安全需求敞口\n",
            "- 安全要求需要进一步细化\n",
            "\n## 五、功能点工作量评估参考\n",
            f"\n根据功能点评估，预估需要额外开发的工作量约为 {analysis_results['workload_estimate']['estimated_person_days']:.0f} 人天\n",
            "\n## 六、风险识别处理\n",
            "\n### 6.1 产品能力与非功能性需求满足度风险\n"
        ])
        
        for risk in analysis_results['consolidated_risks']:
            severity = risk['severity']
            report_parts.append(f"- **{severity.upper()}风险**: {risk['description']}\n")
        
        report_parts.extend([
            "\n## 七、四级风险清单\n",
            "\n| 风险事项 | 风险等级 | 风险状态 | 预计发生时间 | 风险责任人 | 风险影响(成本) | 风险影响(进度) | 风险影响(质量) | 风险触发阈值 | 风险描述 | 风险展朌 | 应对措施 | 风险跟踪 |\n",
            "|----------|----------|----------|--------------|------------|----------------|----------------|----------------|-------------|----------|----------|----------|----------|\n"
        ])
        
        for risk in analysis_results['four_level_risks'][:5]:  # Limit to first 5
            report_parts.append(
                f"| {risk['风险事项']} | {risk['风险等级']} | {risk['风险状态']} | {risk['预计发生时间']} | {risk['风险责任人']} | {risk['风险影响(成本)']} | {risk['风险影响(进度)']} | {risk['风险影响(质量)']} | {risk['风险触发阈值']} | {risk['风险描述'][:30]}... | {risk['风险展朌']} | {risk['应对措施']} | {risk['风险跟踪']} |\n"
            )
        
        report_parts.extend([
            "\n## 八、总结与建议\n",
            f"\n1. **匹配度评估**: 现有产品可满足约{analysis_results['matching_analysis']['match_percentage']:.0f}%的功能需求，剩余{100-int(analysis_results['matching_analysis']['match_percentage']):.0f}%需要定制开发\n",
            "2. **重点关注**: 性能指标、安全要求、集成复杂度需重点评估\n",
            "3. **风险管控**: 建议建立专项风险管理机制，定期跟踪高风险项\n",
            "4. **资源投入**: 建议根据工作量评估合理配置人力资源\n",
            "5. **时间规划**: 建议在原计划基础上考虑风险缓冲时间\n"
        ])
        
        return "".join(report_parts)
    
    def generate_report_doc(self, analysis_results: Dict, title: str = "合同需求评估报告") -> Optional[str]:
        """Generate formatted Word document report from analysis results"""
        if not DOCX_AVAILABLE:
            print("DOCX output requires python-docx. Please install with 'pip install python-docx'")
            return None
        
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.name = '微软雅黑'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add table of contents placeholder
        doc.add_heading('一、需求拆解清单', level=1)
        doc.add_heading('1.1 功能需求', level=2)
        
        # Add functional requirements
        func_reqs = analysis_results['requirements']['functional']
        for req in func_reqs[:10]:  # Limit to first 10 for readability
            doc.add_paragraph(f"{req}功能: 基于合同条款识别", style='List Bullet')
        
        doc.add_heading('1.2 集成需求', level=2)
        doc.add_paragraph('与现有系统的集成对接需求', style='List Bullet')
        doc.add_paragraph('支持多种技术架构', style='List Bullet')
        
        doc.add_heading('1.3 非功能需求', level=2)
        doc.add_heading('性能要求', level=3)
        doc.add_paragraph('系统性能指标需进一步明确')
        
        doc.add_heading('安全要求', level=3)
        doc.add_paragraph('数据安全和权限管理要求')
        
        doc.add_heading('高可用要求', level=3)
        doc.add_paragraph('系统可用性要求')
        
        doc.add_page_break()
        
        doc.add_heading('二、产品功能匹配度分析', level=1)
        doc.add_heading(f'2.1 与现有产品的匹配情况', level=2)
        doc.add_paragraph(f'根据产品功能清单分析，合同中产品功能匹配度约为 {analysis_results["matching_analysis"]["match_percentage"]:.0f}%')
        
        doc.add_heading('2.2 不满足的需求', level=2)
        
        # Add detailed unmatched features list
        unmatched = analysis_results['matching_analysis']['unmatched_features']
        if unmatched:
            doc.add_heading('完全不满足的需求（与产品功能清单匹配不满足内容清单）', level=3)
            for i, item in enumerate(unmatched, 1):
                doc.add_paragraph(f'{i}. {item}: 现有产品暂不支持此功能，需要定制开发', style='List Number')
        
        # Add matched features for transparency
        matched_details = analysis_results['matching_analysis']['matched_features']
        if matched_details:
            doc.add_heading('满足的需求（与产品功能清单匹配内容）', level=3)
            for match_info in matched_details:
                req = match_info['requirement']
                para = doc.add_paragraph(f'{req}: 与以下产品功能匹配', style='List Bullet')
                for feature in match_info['matched_features'][:2]:  # Limit to first 2 for readability
                    module_name = feature.get('模块名称', '未知')
                    feature_name = feature.get('功能清单', '未知')
                    doc.add_paragraph(f'模块: {module_name}, 功能: {feature_name}', style='List Bullet 2')
        
        doc.add_heading('三、技术风险识别', level=1)
        doc.add_heading('3.1 不满足的点', level=2)
        
        tech_risks = analysis_results['technical_risks']
        for risk in tech_risks:
            doc.add_paragraph(f"{risk['type']}: {risk['description']}", style='List Bullet')
        
        doc.add_heading('四、非功能需求拆解', level=1)
        doc.add_heading('4.1 性能需求敞口', level=2)
        doc.add_paragraph('性能指标需要进一步明确')
        
        doc.add_heading('4.2 安全需求敞口', level=2)
        doc.add_paragraph('安全要求需要进一步细化')
        
        doc.add_heading('五、功能点工作量评估参考', level=1)
        doc.add_paragraph(f'根据功能点评估，预估需要额外开发的工作量约为 {analysis_results["workload_estimate"]["estimated_person_days"]:.0f} 人天')
        
        doc.add_heading('六、风险识别处理', level=1)
        doc.add_heading('6.1 产品能力与非功能性需求满足度风险', level=2)
        
        for risk in analysis_results['consolidated_risks']:
            severity = risk['severity']
            doc.add_paragraph(f"{severity.upper()}风险: {risk['description']}", style='List Bullet')
        
        doc.add_heading('七、四级风险清单', level=1)
        
        # Create table for risks
        table = doc.add_table(rows=1, cols=13)
        table.style = 'Table Grid'
        
        # Set header
        hdr_cells = table.rows[0].cells
        headers = ['风险事项', '风险等级', '风险状态', '预计发生时间', '风险责任人', '风险影响(成本)', '风险影响(进度)', '风险影响(质量)', '风险触发阈值', '风险描述', '风险展朌', '应对措施', '风险跟踪']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Add risk data
        for risk in analysis_results['four_level_risks'][:10]:  # Limit to first 10
            row_cells = table.add_row().cells
            row_cells[0].text = risk['风险事项'][:30]  # Truncate if too long
            row_cells[1].text = risk['风险等级']
            row_cells[2].text = risk['风险状态']
            row_cells[3].text = risk['预计发生时间']
            row_cells[4].text = risk['风险责任人']
            row_cells[5].text = risk['风险影响(成本)']
            row_cells[6].text = risk['风险影响(进度)']
            row_cells[7].text = risk['风险影响(质量)']
            row_cells[8].text = risk['风险触发阈值']
            row_cells[9].text = risk['风险描述'][:50]  # Truncate
            row_cells[10].text = risk['风险展朌'][:50]  # Truncate
            row_cells[11].text = risk['应对措施'][:50]  # Truncate
            row_cells[12].text = risk['风险跟踪']
        
        doc.add_heading('八、总结与建议', level=1)
        doc.add_paragraph(f'1. 匹配度评估: 现有产品可满足约{analysis_results["matching_analysis"]["match_percentage"]:.0f}%的功能需求，剩余{100-int(analysis_results["matching_analysis"]["match_percentage"]):.0f}%需要定制开发')
        doc.add_paragraph('2. 重点关注: 性能指标、安全要求、集成复杂度需重点评估')
        doc.add_paragraph('3. 风险管控: 建议建立专项风险管理机制，定期跟踪高风险项')
        doc.add_paragraph('4. 资源投入: 建议根据工作量评估合理配置人力资源')
        doc.add_paragraph('5. 时间规划: 建议在原计划基础上考虑风险缓冲时间')
        
        return doc


def main():
    """Main function to demonstrate the skill"""
    if len(sys.argv) < 3:
        print("Usage: python wenxin5_skill.py <contract_file_path> <traindata_path> [output_format]")
        print("output_format: 'md' for markdown (default), 'doc' for Word document")
        sys.exit(1)
    
    contract_file_path = sys.argv[1]
    traindata_path = sys.argv[2]
    output_format = sys.argv[3] if len(sys.argv) > 3 else 'md'
    
    # Read contract content
    try:
        with open(contract_file_path, 'r', encoding='utf-8') as f:
            contract_content = f.read()
    except UnicodeDecodeError:
        # Try different encoding
        try:
            with open(contract_file_path, 'r', encoding='gbk') as f:
                contract_content = f.read()
        except UnicodeDecodeError:
            # Try utf-8 with error handling
            with open(contract_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                contract_content = f.read()
    
    # Initialize the skill
    skill = ContractAuditSkill(traindata_path)
    
    # Analyze the contract
    results = skill.analyze_contract(contract_content)
    
    # Generate report based on format
    if output_format.lower() == 'doc':
        doc_report = skill.generate_report_doc(results)
        if doc_report is not None:
            # Save as Word document
            output_path = contract_file_path.replace('.txt', '_analysis_report.docx').replace('.docx', '_analysis_report.docx').replace('.pdf', '_analysis_report.docx').replace('.md', '_analysis_report.docx')
            doc_report.save(output_path)
            print(f"Word document report saved to: {output_path}")
        else:
            print("Failed to generate Word document report. Please install python-docx.")
            # Fallback to markdown
            report = skill.generate_report_md(results)
            output_path = contract_file_path.replace('.txt', '_analysis_report.md').replace('.docx', '_analysis_report.md').replace('.pdf', '_analysis_report.md')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report)
            print(f"Fallback markdown report saved to: {output_path}")
    else:
        # Generate markdown report
        report = skill.generate_report_md(results)
        
        # Print report
        print(report)
        
        # Save to file
        output_path = contract_file_path.replace('.txt', '_analysis_report.md').replace('.docx', '_analysis_report.md').replace('.pdf', '_analysis_report.md')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)
        print(f"\nMarkdown report saved to: {output_path}")


if __name__ == "__main__":
    main()